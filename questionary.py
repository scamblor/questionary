#
#  Program to convert questionnaries in YAML format to several other formats
#  like CSV spreadsheet, Moodle XML, json for JavaScript or Docx document.
#
#  Questionary Copyright (c) 2021 Carlos Pardo
#
#  This program is free software: you can redistribute it and/or modify
#  it under the terms of the GNU General Public License as published by
#  the Free Software Foundation, either version 3 of the License, or
#  (at your option) any later version.
#
#  This program is distributed in the hope that it will be useful,
#  but WITHOUT ANY WARRANTY; without even the implied warranty of
#  MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#  GNU General Public License for more details.
#
#  You should have received a copy of the GNU General Public License
#  along with this program.  If not, see <https://www.gnu.org/licenses/>.
#

import re
import os
import codecs
import base64
import random
import copy
import hashlib
import shutil

import yaml
import jinja2
from PIL import Image
import docx
from docx.shared import Inches, Cm, Pt


class Questionary():

   def __init__(self):
      self.templates_path = 'templates'
      self.px_cm = 59    # Pixel per centimeter. Used in docx images. 59 px_cm = 150dpi
      self.hash_len = 20
      self.csv_delimiter = ','
      self.questions = []
      random.seed(1000)
      
      
   def read_yaml(self, filename):
      """Read questions from Yaml file"""
      with codecs.open(filename, 'r', encoding='utf-8') as yamlfile:
         yamldata = yamlfile.read()
      self.questions = yaml.safe_load(yamldata)
      self.yaml_file = filename
      self.filename = os.path.splitext(os.path.basename(filename))[0]
      self.add_image_info()
      self.add_title()
      print('   Readed %d questions' % len(self.questions))


   def add_title(self):
      """Add Title to every question if does not exists"""
      for question in self.questions:
         if not 'Title' in question or not question['Title']:
            question['Title'] = question['Question']


   def text_safe(self, text):
      """Transforma un texto en minúsculas sin acentos y sin espacios"""
      text = text.lower().strip()
      replaceitems = [
         ['[ \t_-]+', '_'],
         ['á', 'a'],
         ['é', 'e'],
         ['í', 'i'],
         ['ó', 'o'],
         ['ú', 'u'],
         ['ö', 'o'],
         ['ü', 'u'],
         ['ñ', 'ni'],
         ]
      for pattern, repl in replaceitems:
         text = re.sub(pattern, repl, text)
      return text
   
   
   def write_yaml(self, filename):
      yamldata = []
      for row in self.questions:
         yamldata.append(
            '- Question: "' + str(row['Question']) + '"\n' +
            '  Image: "' + str(row['Image']) + '"\n' +
            '  Choices:\n'
            '    - "' + str(row['Choices'][0]) + '"\n' +
            '    - "' + str(row['Choices'][1]) + '"\n' +
            '    - "' + str(row['Choices'][2]) + '"\n' +
            '    - "' + str(row['Choices'][3]) + '"\n' +
            '  Block: "' + str(row['Block']) + '"\n'
            )
      with codecs.open(filename, 'w', encoding='utf-8') as yamlfile:
         yamlfile.write('\n'.join(yamldata))
   
   
   def read_csv(self, filename):
      with codecs.open(filename, encoding='utf-8') as csvfile:
         csv_questions = [question for question in csv.reader(csvfile, delimiter=self.csv_delimiter)]
      questions = []
      for question in csv_questions:
         questions.append( {
            'Question': question[0],
            'Image': question[1],
            'Choices': question[2:6],
            'Block': question[6],
            } )
      return questions
   
   
   def write_csv(self, path='./'):
      csv_filename = os.path.join(path, self.filename + '.csv')
      if self.file_newer(csv_filename):
         return
      print('   Writing: ' + csv_filename)

      csv_data = ['Question;Image;Choice_1;Choice_2;Choice_3;Choice_4;Block']
      for row in self.questions:
         line = '"' + row['Question'] + '";"'
         if row['Image']:
            line = line + row['Image']['filename'] + '";"'
         else:
            line = line + '";"'
         line = line + str(row['Choices'][0]) + '";"'
         line = line + str(row['Choices'][1]) + '";"'
         line = line + str(row['Choices'][2]) + '";"'
         line = line + str(row['Choices'][3]) + '";"'
         line = line + self.filename + '"'
         csv_data.append(line)
      with codecs.open(csv_filename, 'w', encoding='utf-8') as csv_file:
         csv_file.write('\n'.join(csv_data))
   
   
   def read_b64(self, filename):
      """Read image and returns data in ascii base64 format"""
      data = open(filename, 'rb').read()
      return base64.b64encode(data).decode('ascii')


   def hashname(self, filename):
      data = open(filename, 'rb').read()
      return hashlib.sha224(data).hexdigest()[:self.hash_len] + os.path.splitext(filename)[1]

   
   def add_image_info(self):
      """Read images from disk and add it several info and translations"""
      for question in self.questions:
         if 'Image' in question and question['Image']:
            imagedict = {}
            imagedict['filename'] = question['Image']
            if os.path.exists(imagedict['filename']):
               imagedict['hashname'] = self.hashname(imagedict['filename'])
               imagedict['path'] = os.path.dirname(imagedict['filename'])
               imagedict['base64'] = self.read_b64(imagedict['filename'])
               width, height = Image.open(imagedict['filename']).size
               imagedict['width'] = width
               imagedict['height'] = height
               if not 'Image_width' in question:
                  imagedict['display_width'] = width
               else:
                  imagedict['display_width'] = question['Image_width']
                  del question['Image_width']
               if imagedict['display_width'] > 800:
                  print('   Warning image file too large. Display width=%3d  File=%s' % (imagedict['display_width'], imagedict['filename']))
               question['Image'] = imagedict                  
            else:
               print('Image file does not exists: ' + imagedict['filename'])
               question['Image'] = {}
         else:
            question['Image'] = {}
   
   
   def jinja_template(self, template_file):
      """Load jinja environment and template file.
         return a jinja template object"""
      templateLoader = jinja2.FileSystemLoader(searchpath=self.templates_path)
      templateEnv = jinja2.Environment(loader=templateLoader)
      self.template = templateEnv.get_template(template_file)
      
   
   def moodle_generate(self, template_file, path='./'):
      """Genera los cuestionarios en formato Moodle xml a partir de las
         cuestiones. Se genera un archivo por cada bloque de cuestiones"""
      xml_filename = os.path.join(path, self.filename + '.xml')
      if self.file_newer(xml_filename):
         return
      print('   Writing: ' + xml_filename)
      self.jinja_template(template_file)
      xml_data = self.template.render(questions = self.questions, category = self.filename)
      with codecs.open(xml_filename, 'w', encoding='utf-8') as outfile:
         outfile.write(xml_data)
   
   
   def docx_make_head(self):
      self.docx = docx.Document()
   
      # Define page properties
      sections = self.docx.sections
      for section in sections:
          section.top_margin = Cm(1)
          section.bottom_margin = Cm(1)
          section.left_margin = Cm(1)
          section.right_margin = Cm(1)
      section = self.docx.sections[0]
      sectPr = section._sectPr
      cols = sectPr.xpath('./w:cols')[0]
      cols.set(docx.oxml.ns.qn('w:num'),'2')
   
      # Define Choice style
      style = self.docx.styles.add_style('Choice', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
      paragraph_format = style.paragraph_format
      paragraph_format.left_indent = Cm(1)
      paragraph_format.first_line_indent = Cm(-0.5)
      paragraph_format.line_spacing = 1
      paragraph_format.space_before = Pt(0)
      paragraph_format.space_after = Pt(0)
      tab_stops = paragraph_format.tab_stops
      tab_stops.add_tab_stop(Cm(4.75))
      tab_stops.add_tab_stop(Cm(5.25))
      paragraph_format.widow_control = True
   
      # Define Image style
      style = self.docx.styles.add_style('Image', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
      paragraph_format = style.paragraph_format
      paragraph_format.line_spacing = 1
      paragraph_format.space_before = Pt(2)
      paragraph_format.space_after = Pt(2)
      paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
   
      # Redefine List Number style
      style = self.docx.styles['List Number']
      paragraph_format = style.paragraph_format
      paragraph_format.space_before = Pt(12)
      paragraph_format.space_after = Pt(0)
      paragraph_format.left_indent = Cm(0.6)
      paragraph_format.first_line_indent = Cm(-0.6)
   
   
   def docx_add_image(self, question):
      """Add image to question in docx document"""
      image_width = question['Image']['display_width']
      par = self.docx.add_paragraph(style='Image')
      par.add_run().add_picture(question['Image']['filename'], width=Cm(image_width / self.px_cm))
   
   
   def suffle_choices(self, question):
       choices = copy.copy(question['Choices'])
       random.shuffle(choices)
       return choices
   
   
   def docx_add_questions(self):
      """Create docx questions"""
      self.docx.add_heading('Bloque de preguntas: ' + self.filename, level=1)
      for question in self.questions:      
         self.docx.add_paragraph(question['Question'], style='List Number')
         if question['Image']:
            self.docx_add_image(question)
         choices = self.suffle_choices(question)
         for i in range(len(choices)):
            self.docx.add_paragraph('abcdefg'[i] + ')\t' + str(choices[i]), style='Choice')
            
   
   def docx_generate(self, path='./'):
      """Genera un archivo docx con las preguntas y opciones de todas
         las cuestiones."""
      docx_filename = os.path.join(path, self.filename + '.docx')
      if self.file_newer(docx_filename):
         return
      print('   Writing: ' + docx_filename)
      self.docx_make_head()
      self.docx_add_questions()
      self.docx.save(docx_filename)


   def file_newer(self, filename):
      if os.path.exists(filename):
         if os.path.getmtime(filename) > os.path.getmtime(self.yaml_file):
            return True
      return False


   def json_generate(self, template_file, path='./'):
      """Genera los archivos json a partir de las cuestiones."""
      # Copy json images
      for question in self.questions:
         if question['Image']:
            dest = os.path.join(path, 'images', question['Image']['hashname'])
            if not os.path.exists(dest):
               shutil.copy2(question['Image']['filename'], dest)

      # Generate json
      json_filename = os.path.join(path, self.filename + '.json')
      if self.file_newer(json_filename):
         return
      print('   Writing: ' + json_filename)
      self.jinja_template(template_file)
      json_data = self.template.render(questions = self.questions)
      with codecs.open(json_filename, 'w', encoding='utf-8') as outfile:
         outfile.write(json_data)


   def html_generate(self, path='./'):
      """Genera una página web con los cuestionarios"""
      self.json_generate(path=path)
   

def main():
   """Main program"""

   moodle_template = 'moodle-multichoice-template.xml'
   json_template = 'json-template.json'

   moodle_path = 'moodle'
   docx_path = 'office'
   csv_path = 'csv'
   html_path = 'docs'
   
   questionary = Questionary()
   
   # Process all yaml files of this directory
   yaml_files = [yaml for yaml in os.listdir('.') if yaml[-5:].lower() == '.yaml']
   for yaml_file in yaml_files:
      if 'license' in yaml_file.lower():
         continue
      print('\nProcessing file: ' + yaml_file)
      questionary.read_yaml(yaml_file)
      questionary.write_csv(path=csv_path)
      questionary.docx_generate(path=docx_path)
      questionary.moodle_generate(moodle_template, path=moodle_path)
      questionary.json_generate(json_template, path=html_path)


main()
